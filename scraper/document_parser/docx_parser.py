"""
DOCX Parser - Extract text from Word documents.
"""
import re
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from loguru import logger


@dataclass
class DocxElement:
    """An element from a DOCX document."""
    type: str  # paragraph, heading, table, image
    text: str
    level: int = 0
    metadata: Dict = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocxParser:
    """
    Extract text, tables, and metadata from DOCX files.
    """

    def __init__(self):
        self.python_docx_available = self._check("docx")

    def _check(self, module: str) -> bool:
        try:
            __import__(module)
            return True
        except ImportError:
            return False

    def parse_file(self, filepath: str) -> Dict[str, Any]:
        """Parse a local DOCX file."""
        if not self.python_docx_available:
            return {"error": "python-docx not installed. Run: pip install python-docx"}

        try:
            from docx import Document

            doc = Document(filepath)
            elements = self._parse_document(doc)

            # Extract metadata
            core_props = self._get_core_properties(filepath)

            return {
                "paragraph_count": len([e for e in elements if e.type == "paragraph"]),
                "heading_count": len([e for e in elements if e.type == "heading"]),
                "table_count": len([e for e in elements if e.type == "table"]),
                "elements": [
                    {"type": e.type, "text": e.text, "level": e.level}
                    for e in elements
                ],
                "full_text": "\n\n".join(e.text for e in elements if e.text),
                "metadata": core_props,
            }

        except Exception as e:
            logger.error(f"DOCX parse failed: {e}")
            return {"error": str(e)}

    def parse_url(self, url: str) -> Dict[str, Any]:
        """Parse DOCX from URL."""
        try:
            import requests
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            from docx import Document
            from io import BytesIO

            doc = Document(BytesIO(response.content))
            elements = self._parse_document(doc)

            return {
                "elements": [
                    {"type": e.type, "text": e.text, "level": e.level}
                    for e in elements
                ],
                "full_text": "\n\n".join(e.text for e in elements if e.text),
            }

        except Exception as e:
            logger.error(f"DOCX URL parse failed: {e}")
            return {"error": str(e)}

    def _parse_document(self, doc) -> List[DocxElement]:
        """Parse document structure."""
        elements = []

        for para in doc.paragraphs:
            element = self._parse_paragraph(para)
            if element:
                elements.append(element)

        # Parse tables
        for table in doc.tables:
            table_element = self._parse_table(table)
            elements.append(table_element)

        return elements

    def _parse_paragraph(self, para) -> Optional[DocxElement]:
        """Parse a paragraph."""
        text = para.text.strip()

        if not text:
            return None

        # Check if heading
        if para.style.name.startswith("Heading"):
            level = 1
            if "Heading 1" in para.style.name:
                level = 1
            elif "Heading 2" in para.style.name:
                level = 2
            elif "Heading 3" in para.style.name:
                level = 3

            return DocxElement(
                type="heading",
                text=text,
                level=level,
            )

        # Regular paragraph
        return DocxElement(
            type="paragraph",
            text=text,
        )

    def _parse_table(self, table) -> DocxElement:
        """Parse a table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        return DocxElement(
            type="table",
            text="\n".join(" | ".join(row) for row in rows),
            metadata={"rows": rows, "row_count": len(rows)},
        )

    def _get_core_properties(self, filepath: str) -> Dict:
        """Get document metadata."""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            import zipfile

            with zipfile.ZipFile(filepath) as z:
                if "docProps/core.xml" in z.namelist():
                    import xml.etree.ElementTree as ET

                    content = z.read("docProps/core.xml")
                    root = ET.fromstring(content)

                    ns = {
                        "dc": "http://purl.org/dc/elements/1.1/",
                        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
                    }

                    props = {}
                    for elem in root:
                        tag = elem.tag.split("}")[-1]  # Remove namespace
                        props[tag] = elem.text

                    return props

        except Exception as e:
            logger.debug(f"Metadata extraction failed: {e}")

        return {}
